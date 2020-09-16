from docx import Document
from docx.shared import Inches
# 打开文档
document = Document()
# 标题
document.add_heading('集群监控', 0)
# 正文
# p = document.add_paragraph('A plain paragraph having some ')
# p.add_run('bold').bold = True
# p.add_run(' and some ')
# p.add_run('italic.').italic = True
# # 标题2
# document.add_heading('1级标题', level=1)
# # 正文2 带样式
# document.add_paragraph('描述', style='Intense Quote')
#
# document.add_paragraph(
#     '无序列表1', style='List Bullet'
# )
# document.add_paragraph(
#     '已排序列表中的第一项', style='List Number'
# )

# document.add_picture('monty-truth.png', width=Inches(1.25))

records = (
    # ("192.168.1.121", '35%', '62%'),
    # ("192.168.1.121", '35%', '62%'),
    # ("192.168.1.121", '35%', '62%')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '节点ip'
hdr_cells[1].text = 'CPU'
hdr_cells[2].text = 'men'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save('demo3.docx')